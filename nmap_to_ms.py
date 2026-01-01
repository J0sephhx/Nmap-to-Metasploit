#!/usr/bin/env python3
"""
Nmap XML to Metasploit RC Script Generator

This script parses Nmap XML output, maps detected services to safe Metasploit
auxiliary modules, and generates a resource script for automated testing.
"""

import argparse
import xml.etree.ElementTree as ET
import yaml
import json
import os
import requests
import sys
from urllib.parse import urlparse
import re
from datetime import datetime
import markdown
import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import subprocess
import threading
import time

class NmapToMetasploitMapper:
    def __init__(self, nmap_xml_file, target_host=None, config_file=None, llm_api_key=None, safety_level='safe'):
        self.nmap_xml_file = nmap_xml_file
        self.target_host = target_host
        self.config_file = config_file or 'service_module_mapping.yaml'
        self.llm_api_key = llm_api_key
        self.safety_level = safety_level
        self.mapping_config = self.load_mapping_config()
        self.nmap_data = self.parse_nmap_xml()
        self.mapped_modules = []
        self.selected_modules = []  # For interactive mode

    def load_mapping_config(self):
        """Load service-to-module mappings from YAML configuration file."""
        if os.path.exists(self.config_file):
            with open(self.config_file, 'r') as f:
                return yaml.safe_load(f)
        else:
            # Default mapping configuration with safety levels
            default_config = {
                'services': {
                    'ssh': [
                        {'module': 'auxiliary/scanner/ssh/ssh_version', 'description': 'Identify SSH version', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/ssh/ssh_login', 'description': 'Brute-force SSH login (safe mode)', 'options': {'STOP_ON_SUCCESS': 'true'}, 'safety_level': 'extended'}
                    ],
                    'http': [
                        {'module': 'auxiliary/scanner/http/http_title', 'description': 'Get HTTP title', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/http/options', 'description': 'HTTP options method', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/http/robots_txt', 'description': 'Check robots.txt file', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/http/dir_scanner', 'description': 'Directory scanner', 'safety_level': 'extended', 'options': {'PATH': '/'}}
                    ],
                    'https': [
                        {'module': 'auxiliary/scanner/http/http_title', 'description': 'Get HTTPS title', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/http/options', 'description': 'HTTPS options method', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/http/dir_scanner', 'description': 'Directory scanner', 'safety_level': 'extended', 'options': {'PATH': '/'}}
                    ],
                    'ftp': [
                        {'module': 'auxiliary/scanner/ftp/ftp_version', 'description': 'Identify FTP version', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/ftp/anonymous', 'description': 'Check for anonymous FTP access', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/ftp/ftp_login', 'description': 'FTP login scanner', 'safety_level': 'bruteforce', 'options': {'STOP_ON_SUCCESS': 'true'}}
                    ],
                    'smb': [
                        {'module': 'auxiliary/scanner/smb/smb_version', 'description': 'Identify SMB version', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/smb/smb_enumusers', 'description': 'Enumerate SMB users (if guest enabled)', 'safety_level': 'extended'},
                        {'module': 'auxiliary/scanner/smb/smb_login', 'description': 'SMB login scanner', 'safety_level': 'bruteforce', 'options': {'STOP_ON_SUCCESS': 'true'}}
                    ],
                    'snmp': [
                        {'module': 'auxiliary/scanner/snmp/snmp_login', 'description': 'Brute-force SNMP community strings (safe mode)', 'safety_level': 'extended'},
                        {'module': 'auxiliary/scanner/snmp/snmp_enum', 'description': 'Enumerate SNMP information', 'safety_level': 'safe'}
                    ],
                    'telnet': [
                        {'module': 'auxiliary/scanner/telnet/telnet_version', 'description': 'Identify Telnet version', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/telnet/telnet_login', 'description': 'Telnet login scanner', 'safety_level': 'bruteforce', 'options': {'STOP_ON_SUCCESS': 'true'}}
                    ],
                    'smtp': [
                        {'module': 'auxiliary/scanner/smtp/smtp_version', 'description': 'Identify SMTP version', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/smtp/smtp_enum', 'description': 'Enumerate SMTP users', 'safety_level': 'extended'}
                    ],
                    'mysql': [
                        {'module': 'auxiliary/scanner/mysql/mysql_version', 'description': 'Identify MySQL version', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/mysql/mysql_login', 'description': 'Brute-force MySQL login (safe mode)', 'safety_level': 'bruteforce', 'options': {'STOP_ON_SUCCESS': 'true'}}
                    ],
                    'mssql': [
                        {'module': 'auxiliary/scanner/mssql/mssql_ping', 'description': 'Ping MSSQL server', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/mssql/mssql_login', 'description': 'Brute-force MSSQL login (safe mode)', 'safety_level': 'bruteforce', 'options': {'STOP_ON_SUCCESS': 'true'}}
                    ],
                    'oracle': [
                        {'module': 'auxiliary/scanner/oracle/tnslsnr_version', 'description': 'Identify Oracle TNS listener version', 'safety_level': 'safe'},
                        {'module': 'auxiliary/scanner/oracle/oracle_login', 'description': 'Brute-force Oracle login (safe mode)', 'safety_level': 'bruteforce', 'options': {'STOP_ON_SUCCESS': 'true'}}
                    ]
                }
            }
            with open(self.config_file, 'w') as f:
                yaml.dump(default_config, f, default_flow_style=False)
            return default_config

    def parse_nmap_xml(self):
        """Parse Nmap XML output and extract hosts, ports, and services."""
        tree = ET.parse(self.nmap_xml_file)
        root = tree.getroot()
        
        hosts_data = []
        
        for host in root.findall('host'):
            # Get IP address
            ip_address = None
            for addr in host.findall('address'):
                if addr.get('addrtype') == 'ipv4':
                    ip_address = addr.get('addr')
                    break
            
            # Skip if no IPv4 address found or doesn't match target if specified
            if not ip_address or (self.target_host and ip_address != self.target_host):
                continue
            
            # Get hostnames
            hostnames = []
            for hostname_elem in host.findall('hostnames/hostname'):
                hostnames.append(hostname_elem.get('name'))
            
            # Get ports and services
            ports_services = []
            for port in host.findall('ports/port'):
                port_id = int(port.get('portid'))
                protocol = port.get('protocol')
                
                state = port.find('state').get('state') if port.find('state') is not None else 'unknown'
                
                service_elem = port.find('service')
                if service_elem is not None:
                    service_name = service_elem.get('name')
                    service_product = service_elem.get('product', '')
                    service_version = service_elem.get('version', '')
                    service_extrainfo = service_elem.get('extrainfo', '')
                else:
                    service_name = 'unknown'
                    service_product = ''
                    service_version = ''
                    service_extrainfo = ''
                
                if state == 'open':
                    ports_services.append({
                        'port': port_id,
                        'protocol': protocol,
                        'service': service_name,
                        'product': service_product,
                        'version': service_version,
                        'extrainfo': service_extrainfo
                    })
            
            hosts_data.append({
                'ip': ip_address,
                'hostnames': hostnames,
                'ports_services': ports_services
            })
        
        return hosts_data

    def map_services_to_modules(self):
        """Map detected services to Metasploit modules based on configuration."""
        mapped_modules = []
        
        for host_data in self.nmap_data:
            host_ip = host_data['ip']
            for service_info in host_data['ports_services']:
                port = service_info['port']
                service = service_info['service']
                
                # Normalize service name for mapping
                normalized_service = service.lower()
                
                # Check for direct service match
                if normalized_service in self.mapping_config['services']:
                    for module_info in self.mapping_config['services'][normalized_service]:
                        # Filter by safety level
                        if self._should_include_module(module_info['safety_level']):
                            # Get LLM suggestions if available
                            llm_suggestions = self.get_llm_suggestions(service, port)
                            mapped_modules.append({
                                'host': host_ip,
                                'port': port,
                                'service': service,
                                'module': module_info['module'],
                                'description': module_info['description'],
                                'options': module_info.get('options', {}),
                                'safety_level': module_info.get('safety_level', 'safe'),
                                'llm_suggestions': llm_suggestions
                            })
                
                # Check for service aliases or partial matches
                for service_name, modules in self.mapping_config['services'].items():
                    if service_name in normalized_service or normalized_service in service_name:
                        for module_info in modules:
                            if self._should_include_module(module_info['safety_level']):
                                llm_suggestions = self.get_llm_suggestions(service, port)
                                mapped_modules.append({
                                    'host': host_ip,
                                    'port': port,
                                    'service': service,
                                    'module': module_info['module'],
                                    'description': module_info['description'],
                                    'options': module_info.get('options', {}),
                                    'safety_level': module_info.get('safety_level', 'safe'),
                                    'llm_suggestions': llm_suggestions
                                })
        
        self.mapped_modules = mapped_modules
        return mapped_modules

    def _should_include_module(self, module_safety_level):
        """Check if module should be included based on safety level."""
        safety_order = {'safe': 0, 'extended': 1, 'bruteforce': 2}
        requested_level = safety_order.get(self.safety_level, 0)
        module_level = safety_order.get(module_safety_level, 0)
        return module_level <= requested_level

    def get_llm_suggestions(self, service, port):
        """Get module suggestions from LLM if API key is provided."""
        if not self.llm_api_key:
            return []
        
        try:
            # Determine provider based on API key format
            if self.llm_api_key.startswith('sk-') and len(self.llm_api_key) > 20:
                provider = 'openai'
            elif 'gemini' in self.llm_api_key.lower():
                provider = 'gemini'
            else:
                provider = 'openai'  # Default
            
            if provider == 'openai':
                headers = {
                    'Authorization': f'Bearer {self.llm_api_key}',
                    'Content-Type': 'application/json'
                }
                
                prompt = f"""
                For a service '{service}' running on port {port}, suggest safe Metasploit auxiliary modules that could be used for enumeration.
                Also provide a brief commentary on why each module is appropriate.
                Return only a JSON array of objects with 'module', 'description', and 'commentary' keys.
                Example: [{{"module": "auxiliary/scanner/service/module_name", "description": "Module description", "commentary": "Why this module is appropriate"}}]
                """
                
                data = {
                    'model': 'gpt-3.5-turbo',
                    'messages': [{'role': 'user', 'content': prompt}],
                    'temperature': 0.1
                }
                
                response = requests.post('https://api.openai.com/v1/chat/completions', headers=headers, json=data)
                
                if response.status_code == 200:
                    result = response.json()
                    content = result['choices'][0]['message']['content']
                    # Clean up response if needed
                    content = re.sub(r'^```json\s*|\s*```$', '', content, flags=re.MULTILINE)
                    suggestions = json.loads(content)
                    return suggestions
                else:
                    print(f"OpenAI API request failed with status {response.status_code}")
                    return []
                    
            elif provider == 'gemini':
                # Gemini API integration
                import google.generativeai as genai
                genai.configure(api_key=self.llm_api_key)
                model = genai.GenerativeModel('gemini-pro')
                
                prompt = f"""
                For a service '{service}' running on port {port}, suggest safe Metasploit auxiliary modules that could be used for enumeration.
                Also provide a brief commentary on why each module is appropriate.
                Return only a JSON array of objects with 'module', 'description', and 'commentary' keys.
                Example: [{{"module": "auxiliary/scanner/service/module_name", "description": "Module description", "commentary": "Why this module is appropriate"}}]
                """
                
                response = model.generate_content(prompt)
                content = response.text
                content = re.sub(r'^```json\s*|\s*```$', '', content, flags=re.MULTILINE)
                suggestions = json.loads(content)
                return suggestions
                
        except Exception as e:
            print(f"Error calling LLM API: {e}")
            return []

    def interactive_review(self):
        """Interactive mode to accept/reject modules."""
        print("=== INTERACTIVE REVIEW MODE ===")
        print("Review and select modules to include in the script:")
        print()
        
        selected_modules = []
        
        for i, module_info in enumerate(self.mapped_modules):
            print(f"{i+1}. Host: {module_info['host']} | Port: {module_info['port']} | Service: {module_info['service']}")
            print(f"   Module: {module_info['module']}")
            print(f"   Description: {module_info['description']}")
            print(f"   Safety Level: {module_info['safety_level']}")
            if module_info.get('llm_suggestions'):
                print(f"   LLM Commentary: {module_info['llm_suggestions']}")
            print()
            
            while True:
                choice = input(f"Include this module? (y/n/skip all/safety level change): ").lower().strip()
                
                if choice == 'y':
                    selected_modules.append(module_info)
                    break
                elif choice == 'n':
                    break
                elif choice == 'skip all':
                    # Skip all remaining modules
                    selected_modules.extend(self.mapped_modules[i+1:])
                    self.selected_modules = selected_modules
                    return selected_modules
                elif choice.startswith('safety'):
                    new_level = input("Enter new safety level (safe/extended/bruteforce): ").lower().strip()
                    if new_level in ['safe', 'extended', 'bruteforce']:
                        self.safety_level = new_level
                        # Re-map modules with new safety level
                        self.map_services_to_modules()
                        return self.interactive_review()  # Restart interactive review
                else:
                    print("Invalid choice. Please enter y, n, skip all, or safety level change.")
        
        self.selected_modules = selected_modules
        return selected_modules

    def dry_run(self):
        """Print planned modules without generating script."""
        modules_to_show = self.selected_modules if self.selected_modules else self.mapped_modules
        
        print("=== DRY RUN MODE ===")
        print("Planned modules based on Nmap XML:")
        print()
        
        for module_info in modules_to_show:
            print(f"Host: {module_info['host']}")
            print(f"  Port: {module_info['port']}")
            print(f"  Service: {module_info['service']}")
            print(f"  Module: {module_info['module']}")
            print(f"  Description: {module_info['description']}")
            print(f"  Safety Level: {module_info['safety_level']}")
            print(f"  Options: {module_info['options']}")
            if module_info.get('llm_suggestions'):
                print(f"  LLM Suggestions: {module_info['llm_suggestions']}")
            print()

    def generate_rc_script(self, output_file, import_xml=False, workspace_name=None):
        """Generate Metasploit resource script."""
        modules_to_use = self.selected_modules if self.selected_modules else self.mapped_modules
        lines = []
        
        # Add header
        lines.append('# Auto-generated Metasploit resource script')
        lines.append('# Generated from Nmap XML: {}'.format(self.nmap_xml_file))
        lines.append('# Target Host: {}'.format(self.target_host or 'All hosts'))
        lines.append('# Safety Level: {}'.format(self.safety_level))
        lines.append('# Generated on: {}'.format(datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
        lines.append('')
        
        # Create workspace if specified
        if workspace_name:
            lines.append('# Create workspace')
            lines.append('workspace -a {}'.format(workspace_name))
            lines.append('')
        
        # Import XML if specified
        if import_xml:
            lines.append('# Import Nmap XML')
            lines.append('db_import {}'.format(self.nmap_xml_file))
            lines.append('')
        
        # Add modules
        for module_info in modules_to_use:
            lines.append('# Host: {} | Port: {} | Service: {} | Safety: {}'.format(
                module_info['host'], module_info['port'], module_info['service'], module_info['safety_level']))
            lines.append('# Description: {}'.format(module_info['description']))
            if module_info.get('llm_suggestions'):
                lines.append('# LLM Commentary: {}'.format(module_info['llm_suggestions']))
            lines.append('use {}'.format(module_info['module']))
            lines.append('set RHOSTS {}'.format(module_info['host']))
            lines.append('set RPORT {}'.format(module_info['port']))
            
            # Add additional options
            for opt_name, opt_value in module_info['options'].items():
                lines.append('set {} {}'.format(opt_name, opt_value))
            
            lines.append('run -j')  # Run in background
            lines.append('')
        
        # Write to file
        with open(output_file, 'w') as f:
            f.write('\n'.join(lines))
        
        return output_file

    def generate_json_report(self, output_file):
        """Generate JSON report of the mapping plan."""
        modules_to_report = self.selected_modules if self.selected_modules else self.mapped_modules
        report = {
            'nmap_xml_file': self.nmap_xml_file,
            'target_host': self.target_host,
            'safety_level': self.safety_level,
            'generated_on': datetime.now().isoformat(),
            'mapping_plan': modules_to_report
        }
        
        with open(output_file, 'w') as f:
            json.dump(report, f, indent=2)
        
        return output_file

    def generate_markdown_report(self, output_file):
        """Generate Markdown report of the mapping plan."""
        modules_to_report = self.selected_modules if self.selected_modules else self.mapped_modules
        
        md_content = f"""# Nmap to Metasploit Mapping Report

**Generated on:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**Nmap XML File:** {self.nmap_xml_file}
**Target Host:** {self.target_host or 'All hosts'}
**Safety Level:** {self.safety_level}

## Mapping Plan

"""
        
        for module_info in modules_to_report:
            md_content += f"""
### Host: {module_info['host']} | Port: {module_info['port']} | Service: {module_info['service']}

- **Module:** {module_info['module']}
- **Description:** {module_info['description']}
- **Safety Level:** {module_info['safety_level']}
- **Options:** {module_info['options']}
"""
            if module_info.get('llm_suggestions'):
                md_content += "- **LLM Suggestions:**\n"
                for suggestion in module_info['llm_suggestions']:
                    md_content += f"  - {suggestion['module']}: {suggestion['commentary']}\n"
            md_content += "\n"
        
        with open(output_file, 'w') as f:
            f.write(md_content)
        
        return output_file

    def generate_docx_report(self, output_file):
        """Generate DOCX report of the mapping plan."""
        modules_to_report = self.selected_modules if self.selected_modules else self.mapped_modules
        
        doc = docx.Document()
        doc.add_heading('Nmap to Metasploit Mapping Report', 0)
        
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Nmap XML File: {self.nmap_xml_file}')
        doc.add_paragraph(f'Target Host: {self.target_host or "All hosts"}')
        doc.add_paragraph(f'Safety Level: {self.safety_level}')
        
        doc.add_heading('Mapping Plan', level=1)
        
        for module_info in modules_to_report:
            doc.add_heading(f'Host: {module_info["host"]} | Port: {module_info["port"]} | Service: {module_info["service"]}', level=2)
            
            doc.add_paragraph(f'Module: {module_info["module"]}', style='List Bullet')
            doc.add_paragraph(f'Description: {module_info["description"]}', style='List Bullet')
            doc.add_paragraph(f'Safety Level: {module_info["safety_level"]}', style='List Bullet')
            doc.add_paragraph(f'Options: {module_info["options"]}', style='List Bullet')
            
            if module_info.get('llm_suggestions'):
                doc.add_paragraph('LLM Suggestions:', style='List Bullet')
                for suggestion in module_info['llm_suggestions']:
                    doc.add_paragraph(f'{suggestion["module"]}: {suggestion["commentary"]}', style='List Bullet 2')
        
        doc.save(output_file)
        return output_file

    def generate_pdf_report(self, output_file):
        """Generate PDF report of the mapping plan."""
        modules_to_report = self.selected_modules if self.selected_modules else self.mapped_modules
        
        c = canvas.Canvas(output_file, pagesize=letter)
        width, height = letter
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, "Nmap to Metasploit Mapping Report")
        
        # Metadata
        c.setFont("Helvetica", 12)
        y_position = height - 80
        c.drawString(50, y_position, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        y_position -= 20
        c.drawString(50, y_position, f"Nmap XML File: {self.nmap_xml_file}")
        y_position -= 20
        c.drawString(50, y_position, f"Target Host: {self.target_host or 'All hosts'}")
        y_position -= 20
        c.drawString(50, y_position, f"Safety Level: {self.safety_level}")
        y_position -= 40
        
        # Mapping Plan
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y_position, "Mapping Plan")
        y_position -= 30
        
        c.setFont("Helvetica", 10)
        for module_info in modules_to_report:
            if y_position < 100:  # New page if needed
                c.showPage()
                y_position = height - 50
                c.setFont("Helvetica-Bold", 14)
                c.drawString(50, y_position, "Mapping Plan (continued)")
                y_position -= 30
                c.setFont("Helvetica", 10)
            
            c.drawString(50, y_position, f"Host: {module_info['host']} | Port: {module_info['port']} | Service: {module_info['service']}")
            y_position -= 15
            c.drawString(70, y_position, f"Module: {module_info['module']}")
            y_position -= 15
            c.drawString(70, y_position, f"Description: {module_info['description']}")
            y_position -= 15
            c.drawString(70, y_position, f"Safety Level: {module_info['safety_level']}")
            y_position -= 15
            c.drawString(70, y_position, f"Options: {module_info['options']}")
            y_position -= 20
            
            if module_info.get('llm_suggestions'):
                c.drawString(70, y_position, "LLM Suggestions:")
                y_position -= 15
                for suggestion in module_info['llm_suggestions']:
                    if y_position < 100:
                        c.showPage()
                        y_position = height - 50
                    c.drawString(90, y_position, f"{suggestion['module']}: {suggestion['commentary']}")
                    y_position -= 15
                y_position -= 5
        
        c.save()
        return output_file

    def setup_docker_lab(self):
        """Generate Docker Compose file for vulnerable lab setup."""
        docker_compose_content = """
version: '3.8'

services:
  vulnerable-ssh:
    image: diogomonica/docker-bench-security
    ports:
      - "2222:22"
    privileged: true
    volumes:
      - /var/run/docker.sock:/var/run/docker.sock

  vulnerable-web:
    image: vulnerables/web-dvwa
    ports:
      - "8080:80"

  vulnerable-ftp:
    image: delfer/alpine-ftp-server
    ports:
      - "21:21"
      - "21000-21010:21000-21010"
    environment:
      - FTP_USER=test
      - FTP_PASS=test
      - PASV_ADDRESS=0.0.0.0
      - PASV_MIN_PORT=21000
      - PASV_MAX_PORT=21010

  vulnerable-smb:
    image: crazymax/samba-krb5-ad
    ports:
      - "139:139"
      - "445:445"
    environment:
      - SMB_NAME=TEST
      - SMB_DOMAIN=TEST.LOCAL
      - SMB_ADMIN_PASS=Password123!
"""
        
        with open('docker-compose.yml', 'w') as f:
            f.write(docker_compose_content)
        
        print("Docker Compose file created: docker-compose.yml")
        print("To start the lab: docker-compose up -d")
        return 'docker-compose.yml'


def main():
    parser = argparse.ArgumentParser(description='Nmap XML to Metasploit RC Script Generator')
    parser.add_argument('-i', '--input', required=True, help='Nmap XML output file (-oX)')
    parser.add_argument('-o', '--output', help='Output RC script file (default: generated.rc)')
    parser.add_argument('-t', '--target', help='Specific target host to process')
    parser.add_argument('-c', '--config', help='Service-to-module mapping configuration file')
    parser.add_argument('-w', '--workspace', help='Metasploit workspace name')
    parser.add_argument('--import-xml', action='store_true', help='Include db_import command in RC script')
    parser.add_argument('--dry-run', action='store_true', help='Show planned modules without generating script')
    parser.add_argument('--llm-api-key', help='API key for LLM integration (OpenAI/Gemini)')
    parser.add_argument('--report', help='Generate JSON report file')
    parser.add_argument('--markdown-report', help='Generate Markdown report file')
    parser.add_argument('--docx-report', help='Generate DOCX report file')
    parser.add_argument('--pdf-report', help='Generate PDF report file')
    parser.add_argument('--interactive', action='store_true', help='Interactive mode to accept/reject modules')
    parser.add_argument('--safety-level', choices=['safe', 'extended', 'bruteforce'], default='safe', 
                       help='Safety level for modules (default: safe)')
    parser.add_argument('--setup-docker-lab', action='store_true', help='Generate Docker Compose file for vulnerable lab')

    args = parser.parse_args()
    
    # Set default output file
    if not args.output:
        args.output = 'generated.rc'
    
    # Initialize mapper
    mapper = NmapToMetasploitMapper(
        nmap_xml_file=args.input,
        target_host=args.target,
        config_file=args.config,
        llm_api_key=args.llm_api_key,
        safety_level=args.safety_level
    )
    
    # Handle Docker lab setup
    if args.setup_docker_lab:
        mapper.setup_docker_lab()
        return
    
    # Map services to modules
    print("Parsing Nmap XML...")
    mapped_modules = mapper.map_services_to_modules()
    print(f"Found {len(mapped_modules)} module mappings")
    
    # Perform interactive review if requested
    if args.interactive:
        selected_modules = mapper.interactive_review()
        print(f"Selected {len(selected_modules)} modules for inclusion")
    
    # Perform dry run if requested
    if args.dry_run:
        mapper.dry_run()
        return
    
    # Generate RC script
    print(f"Generating RC script: {args.output}")
    rc_file = mapper.generate_rc_script(
        output_file=args.output,
        import_xml=args.import_xml,
        workspace_name=args.workspace
    )
    print(f"RC script generated: {rc_file}")
    
    # Generate JSON report if requested
    if args.report:
        print(f"Generating JSON report: {args.report}")
        report_file = mapper.generate_json_report(output_file=args.report)
        print(f"JSON report generated: {report_file}")
    
    # Generate Markdown report if requested
    if args.markdown_report:
        print(f"Generating Markdown report: {args.markdown_report}")
        report_file = mapper.generate_markdown_report(output_file=args.markdown_report)
        print(f"Markdown report generated: {report_file}")
    
    # Generate DOCX report if requested
    if args.docx_report:
        print(f"Generating DOCX report: {args.docx_report}")
        report_file = mapper.generate_docx_report(output_file=args.docx_report)
        print(f"DOCX report generated: {report_file}")
    
    # Generate PDF report if requested
    if args.pdf_report:
        print(f"Generating PDF report: {args.pdf_report}")
        report_file = mapper.generate_pdf_report(output_file=args.pdf_report)
        print(f"PDF report generated: {report_file}")
    
    print("Done!")


if __name__ == '__main__':
    main()
